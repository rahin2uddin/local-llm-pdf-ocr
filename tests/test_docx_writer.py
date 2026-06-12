from __future__ import annotations

import io

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from local_deepl.api.routers import ocr
from local_deepl.core.docx_writer import convert_markdown_to_docx


def test_convert_markdown_to_docx_structure():
    markdown = """# Main Title
## Page 1
This is a paragraph with **bold text** and *italic text* and `inline code`.

- Bullet item 1
- Bullet item 2

1. Numbered item 1
2. Numbered item 2

## Page 2
Another paragraph.
"""
    stream = convert_markdown_to_docx(markdown)
    assert isinstance(stream, io.BytesIO)

    # Read back document structure using python-docx
    doc = Document(stream)

    # Check paragraphs count
    assert len(doc.paragraphs) > 0

    # Check text content of headings/paragraphs
    texts = [p.text for p in doc.paragraphs]
    assert "Main Title" in texts
    assert "Page 1" in texts
    assert "Bullet item 1" in texts
    assert "Numbered item 1" in texts
    assert "Page 2" in texts
    assert "Another paragraph." in texts

    # Verify inline formatting in the first body paragraph
    para = [p for p in doc.paragraphs if "bold text" in p.text][0]
    # Check runs
    runs = para.runs
    bold_runs = [r for r in runs if r.bold]
    assert len(bold_runs) > 0
    assert bold_runs[0].text == "bold text"

    italic_runs = [r for r in runs if r.italic]
    assert len(italic_runs) > 0
    assert italic_runs[0].text == "italic text"

    code_runs = [r for r in runs if r.font.name == "Courier New"]
    assert len(code_runs) > 0
    assert code_runs[0].text == "inline code"


def test_export_docx_endpoint():
    app = FastAPI()
    app.include_router(ocr.router)
    from local_deepl.api.routers import artifacts
    app.include_router(artifacts.router)
    client = TestClient(app)

    response = client.post(
        "/api/export/docx",
        json={"text": "# Test Document\nHello world"},
    )

    assert response.status_code == 200
    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in response.headers["content-disposition"]

    # Verify we can load the returned bytes as a docx
    docx_bytes = io.BytesIO(response.content)
    doc = Document(docx_bytes)
    assert doc.paragraphs[0].text == "Test Document"
    assert doc.paragraphs[1].text == "Hello world"
